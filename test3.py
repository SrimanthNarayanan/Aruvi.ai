import streamlit as st
import pandas as pd
import re
import google.generativeai as genai
import time
import os
from collections import defaultdict
from io import BytesIO

# HTML & PDF Generation
import markdown
from xhtml2pdf import pisa

# Database imports
import pyodbc
import mysql.connector
import psycopg2
import snowflake.connector

# Visualization
import matplotlib.pyplot as plt
import seaborn as sns
import warnings

# Email
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import tempfile

warnings.filterwarnings("ignore", category=UserWarning, module="pandas")

# --- Page Configuration ---
st.set_page_config(
    page_title="AnalytixHub AI",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Professional Look ---
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 1rem;
    }
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        border: 1px solid #e2e8f0;
        margin: 1rem 0;
    }
    .success-box {
        background: #d1fae5;
        color: #065f46;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #10b981;
    }
    .stButton button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 8px;
        font-weight: 600;
        width: 100%;
    }
    .upload-area {
        border: 2px dashed #cbd5e1;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        background: #f8fafc;
        margin: 1rem 0;
    }
    .mode-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        cursor: pointer;
        transition: transform 0.3s ease;
        margin: 1rem 0;
    }
    .mode-card:hover {
        
    }
    .vector-stats {
        background: linear-gradient(135deg, #a78bfa 0%, #7e22ce 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# --- Configuration ---
GEMINI_API_KEY = "AIzaSyDy1mZ6wMvHZdlW5FCiZ_02xf_5261ZmZ8"

try:
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Failed to configure Gemini API: {e}")


import base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import requests # Needed for making API calls

 

# New Configuration for Modern Authentication (Graph API)
 
OUTLOOK_CONFIG = {
    # SMTP is generally avoided for modern OAuth 2.0/Graph API sending
    "server": "smtp.office365.com", 
    "port": 587,
    
    # Modern Auth Details (needed for Graph API)
    "sender_email": "Srimanth.Narayanan@analytixhub.ai", # Set your email
    "CLIENT_ID": "d8ed44c2-936f-426d-a071-dccd0647ff1d",
    "TENANT_ID": "03a187b1-7692-4998-b9e9-8e8f11252566",
    "CLIENT_SECRET": "nxO8Q~7FTqY5VeMBjGY6BP52KFGpMK8EfJ17UaN-",
    
    # MS Graph API Endpoints
    "TOKEN_URL": f"https://login.microsoftonline.com/03a187b1-7692-4998-b9e9-8e8f11252566/oauth2/v2.0/token",
    "GRAPH_API_URL": "https://graph.microsoft.com/v1.0"
}

# --- Helper Function for Graph API Token Acquisition and Email Sending ---

def _get_access_token():
    """Acquires an access token using Client Credentials Flow."""
    try:
        token_data = {
            'client_id': OUTLOOK_CONFIG['CLIENT_ID'],
            'scope': 'https://graph.microsoft.com/.default', # Permissions scope
            'client_secret': OUTLOOK_CONFIG['CLIENT_SECRET'],
            'grant_type': 'client_credentials',
        }
        # For sending as the user, you'd typically need delegated permissions 
        # and a more complex flow (e.g., Authorization Code Flow)
        
        response = requests.post(OUTLOOK_CONFIG['TOKEN_URL'], data=token_data)
        response.raise_for_status() # Raise exception for bad status codes
        return response.json().get('access_token')
    except Exception as e:
        print(f"Token acquisition failed: {e}")
        return None

def _send_email_via_graph_api(access_token, sender_email, recipient_email, subject, message, attachments):
    """Sends email using Microsoft Graph API 'sendMail' endpoint."""
    if not access_token:
        return False

    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    mail_body = {
        "message": {
            "subject": subject,
            "body": {
                "contentType": "Html",
                "content": message
            },
            "toRecipients": [
                {
                    "emailAddress": {
                        "address": recipient_email
                    }
                }
            ],
            "attachments": attachments
        },
        "saveToSentItems": "true" # Save a copy to the Sent Items folder
    }

    # Use the /sendMail endpoint
    send_mail_url = f"{OUTLOOK_CONFIG['GRAPH_API_URL']}/users/{sender_email}/sendMail"

    try:
        response = requests.post(send_mail_url, headers=headers, json=mail_body)
        response.raise_for_status()
        return True
    except requests.exceptions.RequestException as e:
        print(f"Graph API email sending failed: {e}")
        return False

# --- Main Sending Function Modified for Modern Authentication ---

def send_analysis_email(recipient_email: str, subject: str, message: str, 
                        pdf_bytes: bytes = None, csv_bytes: bytes = None, 
                        excel_bytes: bytes = None) -> bool:
    """
    Send analysis report via email with attachments using Microsoft Graph API (OAuth 2.0).
    """
    try:
        # 1. Prepare attachments for Graph API format
        attachments = []
        
        if pdf_bytes:
            attachments.append({
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": "analysis_report.pdf",
                "contentType": "application/pdf", # Specific content type is better
                "contentBytes": base64.b64encode(pdf_bytes).decode('utf-8')
            })
            
        if csv_bytes:
            attachments.append({
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": "analysis_data.csv",
                "contentType": "text/csv", 
                "contentBytes": base64.b64encode(csv_bytes).decode('utf-8')
            })
            
        if excel_bytes:
            # .xlsx MIME type is application/vnd.openxmlformats-officedocument.spreadsheetml.sheet
            attachments.append({ 
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": "analysis_data.xlsx",
                "contentType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                "contentBytes": base64.b64encode(excel_bytes).decode('utf-8')
            })
            
        # 2. Get the Access Token
        access_token = _get_access_token()
        if not access_token:
            return False

        # 3. Send the email via Graph API
        return _send_email_via_graph_api(
            access_token=access_token,
            sender_email=OUTLOOK_CONFIG["sender_email"],
            recipient_email=recipient_email,
            subject=subject,
            message=message,
            attachments=attachments
        )
        
    except Exception as e:
        print(f"Email sending failed: {str(e)}") # Use print instead of st.error if 'st' is not defined
        return False

# --------------------------------------------------------------------------------
# --- Vector Database RAG System ---
# --------------------------------------------------------------------------------

class VectorRAGSystem:
    def __init__(self, persist_directory="./chroma_db"):
        self.persist_directory = persist_directory
        self.vector_store = None
        self.embeddings = None
        self.text_splitter = None
        self.initialize_components()
    
    def initialize_components(self):
        """Initialize embeddings and text splitter"""
        try:
            # Initialize embeddings
            from sentence_transformers import SentenceTransformer
            self.embeddings = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Initialize text splitter
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
            
            # Initialize vector store
            self.initialize_vector_store()
            
            
            
        except Exception as e:
            st.error(f"❌ Failed to initialize RAG system: {e}")
    
    def initialize_vector_store(self):
        """Initialize ChromaDB vector store"""
        try:
            import chromadb
            from chromadb.config import Settings
            
            # Create client
            self.client = chromadb.PersistentClient(path=self.persist_directory)
            
            # Get or create collection
            try:
                self.collection = self.client.get_collection("document_chunks")
            except:
                self.collection = self.client.create_collection(
                    name="document_chunks",
                    metadata={"description": "Document chunks for RAG system"}
                )
                
        except Exception as e:
            st.error(f"❌ Failed to initialize vector store: {e}")
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF file"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            st.error(f"❌ PDF text extraction failed: {e}")
            return None
    
    def extract_text_from_txt(self, txt_file):
        """Extract text from TXT file"""
        try:
            return str(txt_file.read(), "utf-8")
        except Exception as e:
            st.error(f"❌ TXT file reading failed: {e}")
            return None
    
    def extract_text_from_csv(self, csv_file):
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(csv_file)
            return df.to_string()
        except Exception as e:
            st.error(f"❌ CSV file reading failed: {e}")
            return None
        
    def extract_text_from_excel(self, excel_file):
        """Extract text from Excel file"""
        try:
            df = pd.read_excel(excel_file)
            text_content = ""
            
            # Add sheet names and content
            if isinstance(df, dict):  # Multiple sheets
                for sheet_name, sheet_data in df.items():
                    text_content += f"--- Sheet: {sheet_name} ---\n"
                    text_content += sheet_data.to_string() + "\n\n"
            else:  # Single sheet
                text_content = df.to_string()
                
            return text_content.strip()
        except Exception as e:
            st.error(f"❌ Excel file reading failed: {e}")
            return None

    def extract_text_from_docx(self, docx_file):
        """Extract text from Word document"""
        try:
            from docx import Document
            doc = Document(docx_file)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " | "
                    text_content += "\n"
                text_content += "\n"
                
            return text_content.strip()
        except Exception as e:
            st.error(f"❌ DOCX file reading failed: {e}")
            return None

    def extract_text_from_json(self, json_file):
        """Extract text from JSON file"""
        try:
            import json
            data = json.load(json_file)
            
            def flatten_json(data, parent_key='', separator='.'):
                """Flatten JSON structure"""
                items = []
                if isinstance(data, dict):
                    for k, v in data.items():
                        new_key = f"{parent_key}{separator}{k}" if parent_key else k
                        if isinstance(v, (dict, list)):
                            items.extend(flatten_json(v, new_key, separator=separator).items())
                        else:
                            items.append((new_key, str(v)))
                elif isinstance(data, list):
                    for i, v in enumerate(data):
                        new_key = f"{parent_key}{separator}{i}" if parent_key else str(i)
                        if isinstance(v, (dict, list)):
                            items.extend(flatten_json(v, new_key, separator=separator).items())
                        else:
                            items.append((new_key, str(v)))
                return dict(items)
            
            flattened = flatten_json(data)
            return "\n".join([f"{k}: {v}" for k, v in flattened.items()])
        except Exception as e:
            st.error(f"❌ JSON file reading failed: {e}")
            return None
    
    def process_document(self, uploaded_file):
        """Process uploaded document and store in vector database"""
        try:
            # Extract text based on file type
            file_extension = uploaded_file.name.split('.')[-1].lower()
            text_content = None
            
            if file_extension == 'pdf':
                text_content = self.extract_text_from_pdf(uploaded_file)
            elif file_extension == 'txt':
                text_content = self.extract_text_from_txt(uploaded_file)
            elif file_extension == 'csv':
                text_content = self.extract_text_from_csv(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                text_content = self.extract_text_from_excel(uploaded_file)
            elif file_extension == 'docx':
                text_content = self.extract_text_from_docx(uploaded_file)
            elif file_extension == 'json':
                text_content = self.extract_text_from_json(uploaded_file)
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_extension}"
                }
            
            if not text_content:
                return {
                    "status": "error",
                    "message": "Could not extract text from document"
                }
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text_content)
            
            # Generate embeddings and store in vector database
            embeddings = self.embeddings.encode(chunks).tolist()
            
            # Prepare documents for storage
            documents = []
            metadatas = []
            ids = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                documents.append(chunk)
                metadatas.append({
                    "filename": uploaded_file.name,
                    "chunk_index": i,
                    "file_type": file_extension,
                    "file_size": uploaded_file.size,
                    "total_chunks": len(chunks)
                })
                ids.append(f"{uploaded_file.name}_chunk_{i}")
            
            # Add to vector store
            self.collection.add(
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            
            return {
                "status": "success",
                "message": f"Successfully processed {uploaded_file.name}",
                "filename": uploaded_file.name,
                "chunks_added": len(chunks),
                "file_type": file_extension,
                "file_size": uploaded_file.size,
                "total_chunks": len(chunks)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Document processing failed: {str(e)}"
            }
    
    def search_documents(self, query: str, n_results: int = 5):
        """Search for relevant document chunks"""
        try:
            # Generate query embedding
            query_embedding = self.embeddings.encode([query]).tolist()[0]
            
            # Search in vector database
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                include=["documents", "metadatas", "distances"]
            )
            
            return results
            
        except Exception as e:
            st.error(f"❌ Vector search failed: {e}")
            return None
    
    def get_relevant_context(self, query: str, n_results: int = 5, max_tokens: int = 4000) -> str:
        """Get relevant context with token control"""
        results = self.search_documents(query, n_results)
        
        if not results or not results['documents']:
            return "No relevant document context found."
        
        context_parts = []
        current_tokens = 0
        
        for i, (doc, metadata, distance) in enumerate(zip(
            results['documents'][0], 
            results['metadatas'][0], 
            results['distances'][0]
        )):
            # Estimate tokens (roughly 4 chars per token)
            doc_tokens = len(doc) // 4
            
            if current_tokens + doc_tokens > max_tokens:
                break
                
            context_parts.append(f"\n--- Excerpt {i+1} (Relevance: {1-distance:.3f}) ---")
            context_parts.append(f"Source: {metadata['filename']} | Chunk {metadata['chunk_index']+1}/{metadata['total_chunks']}")
            context_parts.append(f"Content: {doc}")
            current_tokens += doc_tokens
        
        if not context_parts:
            return "Relevant context found but too large to process. Try a more specific question."
        
        return "\n".join(context_parts)
    
    
    
    def get_collection_stats(self):
        """Get statistics about the vector database"""
        try:
            count = self.collection.count()
            return {
                "total_chunks": count,
                "collection_name": "document_chunks",
                "persist_directory": self.persist_directory
            }
        except Exception as e:
            return {"error": str(e)}
    
    def clear_documents(self):
        """Clear all documents from vector database by dropping and recreating the collection."""
        try:
            collection_name = "document_chunks"
            
            # 1. Delete the existing collection
            self.client.delete_collection(collection_name)
            
            # 2. Recreate the collection immediately (This resets the database)
            self.collection = self.client.create_collection(
                name=collection_name,
                metadata={"description": "Document chunks for RAG system"}
            )
            return True
        except Exception as e:
            st.error(f"❌ Failed to clear documents: {e}")
            return False

# Initialize RAG system
rag_system = VectorRAGSystem()

# --------------------------------------------------------------------------------
# --- Core Database Functions (Keep your existing ones) ---
# --------------------------------------------------------------------------------

def connect_to_database(db_type, credentials):
    """Direct database connection"""
    try:
        if db_type == "SQL Server":
            conn = pyodbc.connect(f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={credentials['host']},{credentials['port']};DATABASE={credentials['dbname']};UID={credentials['user']};PWD={credentials['password']}")
        elif db_type == "MySQL":
            conn = mysql.connector.connect(
                host=credentials['host'],
                port=credentials['port'],
                database=credentials['dbname'],
                user=credentials['user'],
                password=credentials['password']
            )
        elif db_type == "PostgreSQL":
            conn = psycopg2.connect(
                host=credentials['host'],
                port=credentials['port'],
                dbname=credentials['dbname'],
                user=credentials['user'],
                password=credentials['password']
            )
        elif db_type == "Snowflake":
            conn = snowflake.connector.connect(
                account=credentials['account'],
                warehouse=credentials['warehouse'],
                database=credentials['database'],
                user=credentials['user'],
                password=credentials['password']
            )
        else:
            return None
        return conn
    except Exception as e:
        st.error(f"❌ Connection failed: {e}")
        return None

def get_db_schema(conn, db_type):
    """Direct schema extraction"""
    try:
        if db_type == "Snowflake": 
            return get_snowflake_schema(conn)
        elif db_type == "SQL Server": 
            return get_sql_server_schema(conn)
        elif db_type == "MySQL": 
            return get_mysql_schema(conn)
        elif db_type == "PostgreSQL": 
            return get_postgresql_schema(conn)
        else:
            return [], []
    except Exception as e:
        st.error(f"Failed to retrieve schema: {e}")
        return [], []

def get_snowflake_schema(conn):
    cursor = conn.cursor()
    db_name = conn.database
    cols_sql = f"SELECT TABLE_SCHEMA, TABLE_NAME, COLUMN_NAME, DATA_TYPE FROM {db_name}.INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA != 'INFORMATION_SCHEMA' ORDER BY TABLE_SCHEMA, TABLE_NAME, ORDINAL_POSITION;"
    cursor.execute(cols_sql)
    cols = [dict(zip([d[0] for d in cursor.description], row)) for row in cursor.fetchall()]
    return cols, []

def get_sql_server_schema(conn):
    tables_query = "SELECT s.name AS TABLE_SCHEMA, t.name AS TABLE_NAME, c.name AS COLUMN_NAME, ty.name AS DATA_TYPE FROM sys.tables t JOIN sys.schemas s ON t.schema_id = s.schema_id JOIN sys.columns c ON t.object_id = c.object_id JOIN sys.types ty ON c.user_type_id = ty.user_type_id ORDER BY s.name, t.name, c.column_id"
    cols_df = pd.read_sql(tables_query, conn)
    return cols_df.to_dict('records'), []

def get_mysql_schema(conn):
    db_name = conn.database
    cols_sql = f"SELECT TABLE_SCHEMA, TABLE_NAME, COLUMN_NAME, DATA_TYPE FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA = '{db_name}' ORDER BY TABLE_NAME, ORDINAL_POSITION;"
    cols_df = pd.read_sql(cols_sql, conn)
    return cols_df.to_dict('records'), []

def get_postgresql_schema(conn):
    cols_sql = "SELECT table_schema as TABLE_SCHEMA, table_name as TABLE_NAME, column_name as COLUMN_NAME, data_type as DATA_TYPE FROM information_schema.columns WHERE table_schema = 'public' ORDER BY table_name, ordinal_position;"
    cols_df = pd.read_sql(cols_sql, conn)
    return cols_df.to_dict('records'), []

# --------------------------------------------------------------------------------
# --- Enhanced AI Functions with RAG ---
# --------------------------------------------------------------------------------

def generate_sql_with_gemini(question: str, schema_card: str, db_type: str, rag_context: str = "") -> str:
    """Direct SQL generation with Gemini and RAG context"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    
    enhanced_prompt = f"""You are an expert {db_type} database analyst. Generate a precise SELECT query. 
STRICT REQUIREMENTS: 
- Return EXACTLY one SELECT statement. 
- Use ONLY SELECT, FROM, JOIN, WHERE, GROUP BY, HAVING, ORDER BY, LIMIT clauses. 
- NO DML or DDL statements. 
- NO trailing semicolons at the end
- Use proper JOINs based on the relationships shown. 
- Do not include any explanations, comments, or markdown formatting
- Return only the pure SQL query

AVAILABLE SCHEMA: 
{schema_card}"""

    if rag_context and rag_context != "No relevant document context found.":
        enhanced_prompt += f"""

BUSINESS CONTEXT FROM DOCUMENTS:
{rag_context}"""

    enhanced_prompt += f"""

BUSINESS QUESTION: "{question}" 

Generate only the SQL SELECT statement (no explanations, no markdown backticks, no semicolons):"""
    
    try:
        response = model.generate_content(enhanced_prompt)
        sql_query = response.text.strip()
        
        # Additional cleaning
        sql_query = re.sub(r'^```sql\s*|\s*```$', '', sql_query).strip()
        sql_query = sql_query.rstrip(';')
        
        return sql_query
    except Exception as e:
        st.error(f"Error during SQL generation: {e}")
        return ""

def generate_visualization_code(df: pd.DataFrame, question: str) -> str:
    """Generate visualization code"""
    if df.empty or len(df.columns) < 2: return "# Not enough data to generate meaningful visualizations."
    model = genai.GenerativeModel("gemini-2.5-flash")
    data_summary = df.head(20).to_csv(index=False)
    prompt = f"""You are a senior data analyst. Create two distinct, insightful visualizations. USER QUESTION: "{question}" DATA SUMMARY (CSV format): {data_summary} Generate Python code for two side-by-side charts using matplotlib/seaborn. Requirements: - Create figure with: fig, axes = plt.subplots(1, 2, figsize=(14, 6)) - Plot on axes[0] and axes[1] - Add titles and labels - Use plt.tight_layout() - DO NOT call plt.show() - Use double quotes (") for all string literals (e.g., titles, labels) in the code. Return ONLY the raw Python code."""
    try:
        response = model.generate_content(prompt)
        raw_code = response.text.strip()
        return re.sub(r"```(?:python)?\n(.*?)```", r"\1", raw_code, flags=re.DOTALL).strip()
    except Exception as e:
        return f"# Visualization code generation failed: {e}"

def generate_insight_with_gemini(df: pd.DataFrame, question: str, rag_context: str = "") -> str:
    """Direct insight generation with RAG context"""
    if df.empty: 
        return "The query returned no data."
    
    model = genai.GenerativeModel("gemini-2.5-flash")
    data_summary = df.head(20).to_csv(index=False)
    
    if len(data_summary) > 4000: 
        data_summary = data_summary[:4000] + "\n... (data truncated)"
    
    prompt = f"""You are a business analyst. Analyze the following data which was generated to answer the question: "{question}" """
    
    if rag_context and rag_context != "No relevant document context found.":
        prompt += f"""

ADDITIONAL BUSINESS CONTEXT FROM DOCUMENTS:
{rag_context}"""
    
    prompt += f"""

DATA TO ANALYZE:
{data_summary}

Provide concise, business-ready insights in markdown format. 
Focus on key trends, outliers, and actionable recommendations. 
Respond in bullet points and use tables if appropriate."""
    
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating insights: {e}")
        return "Failed to generate insights."

def generate_document_analysis_with_rag(question: str, rag_context: str = "") -> str:
    """Analyze documents using RAG system"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    
    # If rag_context is not provided, get it from the vector database
    if not rag_context:
        rag_context = rag_system.get_relevant_context(question, n_results=5)
    
    prompt = f"""You are a business analyst. Analyze the relevant document excerpts to answer the user's question.

USER QUESTION: {question}

RELEVANT DOCUMENT EXCERPTS:
{rag_context}

Based on the provided document excerpts, provide a comprehensive analysis that includes:

## 📊 Key Findings
- Main insights and discoveries from the documents
- Important patterns or trends identified

## 📈 Data Points & Statistics  
- Relevant numbers, metrics, or statistics mentioned
- Quantitative information found

## 💼 Business Implications
- How this information affects business decisions
- Opportunities or risks identified

## 🎯 Recommendations
- Actionable suggestions based on the analysis
- Next steps or areas for further investigation

## ⚠️ Limitations & Gaps
- Any missing information or context gaps
- Assumptions made in the analysis

Provide your analysis in clear, business-friendly markdown format with appropriate headings and bullet points."""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"## ❌ Analysis Error\n\nFailed to analyze documents: {str(e)}"


# --------------------------------------------------------------------------------
# --- Session Management Functions ---
# --------------------------------------------------------------------------------

def manage_session_persistence():
    """Manage session state to prevent unwanted reruns"""
    # List of keys that should persist across reruns
    persistent_keys = [
        'logged_in', 'app_mode', 'db_conn', 'db_type', 'sql_query', 
        'query_result', 'insight', 'viz_code', 'last_question', 'active_tab',
        'uploaded_documents', 'doc_analysis_results', 'email_sent',
        'recipient_email', 'email_message', 'email_format', 'rag_clear_success',
        'f_db_type', 'f_account', 'f_warehouse', 'f_database', 'f_schema', 
        'f_user_sf', 'f_password_sf', 'f_host', 'f_port', 'f_dbname', 
        'f_user_gen', 'f_password_gen'
    ]
    
    # Ensure all persistent keys exist in session state
    for key in persistent_keys:
        if key not in st.session_state:
            if key == 'uploaded_documents':
                st.session_state[key] = []
            elif key == 'doc_analysis_results':
                st.session_state[key] = {'analysis': '', 'rag_context': '', 'question': '', 'analysis_type': ''}
            elif key in ['email_sent', 'rag_clear_success']:
                st.session_state[key] = False
            elif key in ['recipient_email', 'email_message']:
                st.session_state[key] = ""
            elif key == 'email_format':
                st.session_state[key] = "PDF Report"
            # Initialize form keys to default values
            elif key == 'f_db_type':
                 st.session_state[key] = "Snowflake"
            elif key == 'f_port':
                 st.session_state[key] = "443"
            elif key in ['f_account', 'f_warehouse', 'f_database', 'f_schema', 
                         'f_user_sf', 'f_password_sf', 'f_host', 'f_dbname', 
                         'f_user_gen', 'f_password_gen']:
                 st.session_state[key] = ""
            elif key == 'f_schema':
                 st.session_state[key] = "PUBLIC"
            else:
                st.session_state[key] = None

def init_session_state():
    """Initializes session state variables if they don't exist."""
    manage_session_persistence() # Ensure persistence first
    # Additional initialization or resetting logic if needed

def reset_email_state():
    """Reset email-related session state"""
    st.session_state.email_sent = False
    st.session_state.recipient_email = ""
    st.session_state.email_message = ""
    st.session_state.email_format = "PDF Report"

def safe_rerun():
    """Safe rerun that preserves essential session state"""
    st.rerun()

def pick_relevant_tables(columns: list[dict], prompt: str, max_tables: int = 10) -> set[str]:
    """Find relevant tables based on prompt keywords"""
    if not prompt: return set()
    words = set(re.findall(r"[A-Za-z0-9_]+", prompt.lower()))
    scores = defaultdict(float)
    for c in columns:
        schema, table = c.get('TABLE_SCHEMA', 'PUBLIC'), c.get('TABLE_NAME', '')
        if not table: continue
        full_table = f"{schema}.{table}"
        if table.lower() in words: scores[full_table] += 10.0
        for word in words:
            if word in table.lower(): scores[full_table] += 5.0
        col_name = c.get('COLUMN_NAME', '').lower()
        if col_name in words: scores[full_table] += 3.0
        else:
            for word in words:
                if word in col_name: scores[full_table] += 1.0
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    top_tables = {t for t, s in ranked if s > 0}
    return set(list(top_tables)[:max_tables])

def build_schema_card(columns: list[dict], fks: list[dict], tables_filter: set[str] | None = None) -> str:
    """Build schema description for AI"""
    table_cols = defaultdict(list)
    for c in columns:
        full_table_name = f"{c.get('TABLE_SCHEMA', 'PUBLIC')}.{c['TABLE_NAME']}"
        if tables_filter and full_table_name not in tables_filter: continue
        table_cols[full_table_name].append(f"{c['COLUMN_NAME']}:{c['DATA_TYPE']}")
    lines = ["Schema (relevant tables and columns):"]
    for table, cols in sorted(table_cols.items()):
        lines.append(f"* {table}: {', '.join(cols)}")
    return "\n".join(lines)

def is_safe_select(sql: str) -> bool:
    """Validate SQL safety"""
    if not isinstance(sql, str) or not sql.strip(): return False
    dangerous_keywords = ['INSERT', 'UPDATE', 'DELETE', 'MERGE', 'CREATE', 'ALTER', 'DROP', 'TRUNCATE', 'GRANT', 'REVOKE', 'EXECUTE', 'DECLARE', 'BEGIN']
    sql_no_comments = re.sub(r'--.*?\n|/\*.*?\*/', ' ', sql, flags=re.DOTALL)
    sql_normalized = ' '.join(sql_no_comments.upper().split())
    if not sql_normalized.startswith("SELECT"): return False
    for keyword in dangerous_keywords:
        if re.search(r'\b' + keyword + r'\b', sql_normalized): return False
    return True

def create_html_pdf(html_content: str, question: str) -> bytes:
    """Generate PDF from HTML content with proper formatting and spacing"""
    def format_content_for_pdf(content):
        """Convert markdown content to properly formatted HTML for PDF"""
        # First convert markdown to HTML
        html_content = markdown.markdown(content, extensions=['tables'])
        # Add proper spacing and styling for PDF
        html_content = html_content.replace('<h1>', '<div class="section"><h1>')
        html_content = html_content.replace('</h1>', '</h1></div>')
        html_content = html_content.replace('<h2>', '<div class="section"><h2>')
        html_content = html_content.replace('</h2>', '</h2></div>')
        html_content = html_content.replace('<h3>', '<div class="subsection"><h3>')
        html_content = html_content.replace('</h3>', '</h3></div>')
        # Replace emojis with text labels
        emoji_replacements = {
            '📊': '<span class="section-icon">KEY FINDINGS</span>', 
            '📈': '<span class="section-icon">DATA POINTS</span>', 
            '💼': '<span class="section-icon">BUSINESS IMPLICATIONS</span>', 
            '🎯': '<span class="section-icon">RECOMMENDATIONS</span>', 
            '⚠️': '<span class="section-icon">LIMITATIONS</span>', 
            '🔍': '<span class="section-icon">ANALYSIS</span>', 
            '💰': '<span class="currency">INR</span>', 
            '•': '•' 
        }
        for emoji, replacement in emoji_replacements.items():
            html_content = html_content.replace(emoji, replacement)
        # Fix currency symbols
        html_content = html_content.replace('â‚¹', '₹')
        html_content = html_content.replace('₹', '<span class="currency">INR</span> ')
        return html_content

    css_style = """ 
    <style> 
    @page { size: a4 portrait; margin: 1.5cm; } 
    body { font-family: "Helvetica", "Arial", sans-serif; font-size: 11pt; line-height: 1.6; color: #333333; margin: 0; padding: 0; } 
    .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 25px; text-align: center; margin-bottom: 30px; border-radius: 8px; } 
    .header h1 { font-size: 24pt; margin: 0; font-weight: bold; } 
    .header h2 { font-size: 16pt; margin: 10px 0 0 0; font-weight: normal; opacity: 0.9; } 
    .metadata { background-color: #f8fafc; padding: 20px; border-radius: 8px; margin-bottom: 25px; border-left: 4px solid #667eea; } 
    .metadata strong { color: #1E3A8A; } 
    .section { margin-bottom: 30px; padding-bottom: 20px; border-bottom: 1px solid #e2e8f0; page-break-inside: avoid; } 
    .subsection { margin-bottom: 15px; padding-left: 10px; } 
    h1 { color: #1E3A8A; font-size: 18pt; margin-bottom: 15px; padding-bottom: 8px; border-bottom: 2px solid #667eea; } 
    h2 { color: #374151; font-size: 14pt; margin: 20px 0 12px 0; background-color: #f1f5f9; padding: 10px 15px; border-radius: 5px; } 
    h3 { color: #4B5563; font-size: 12pt; margin: 15px 0 8px 0; } 
    .section-icon { background-color: #667eea; color: white; padding: 4px 8px; border-radius: 4px; font-size: 9pt; font-weight: bold; margin-right: 8px; } 
    p { margin: 10px 0; text-align: justify; } 
    ul, ol { margin: 12px 0; }
    table { width: 100%; border-collapse: collapse; margin-top: 15px; } 
    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; } 
    th { background-color: #DBEAFE; font-weight: bold; color: #1E3A8A; }
    </style> """ 

    full_html = f"""
    <html>
    <head>{css_style}</head>
    <body>
        <div class="header">
            <h1>Analytix Hub Report</h1>
            <h2>AI-Powered Business Intelligence</h2>
        </div>
        <div class="metadata">
            <strong>Query:</strong> "{question}"<br>
            <strong>Date:</strong> {time.strftime('%Y-%m-%d %H:%M:%S')}
        </div>
        {format_content_for_pdf(html_content)}
    </body>
    </html>
    """
    result = BytesIO()
    pdf = pisa.CreatePDF(BytesIO(full_html.encode("UTF-8")), dest=result)
    if not pdf.err:
        return result.getvalue()
    else:
        st.error("Error creating PDF report.")
        return b""

def reset_analysis_state():
    """Resets the state of the analysis outputs."""
    st.session_state.sql_query = None
    st.session_state.query_result = None
    st.session_state.insight = None
    st.session_state.viz_code = None
    st.session_state.email_sent = False

def clear_db_callback():
    """Callback function for clearing vector DB."""
    if rag_system.clear_documents():
        st.session_state.rag_clear_success = True
        st.session_state.doc_analysis_results = {'analysis': '', 'rag_context': '', 'question': '', 'analysis_type': ''}
        st.session_state.uploaded_documents = []


# --------------------------------------------------------------------------------
# --- UI Display Functions (Changes applied here) ---
# --------------------------------------------------------------------------------

def display_mode_selection():
    st.markdown('<h1 class="main-header">Welcome to Clara.ai 🧠</h1>', unsafe_allow_html=True)
    st.subheader("Select your mode of analysis:")

    col1, col2 = st.columns(2)

    with col1:
        # Retain custom style with explicit button for reliability
        st.markdown('<div class="mode-card" style="margin-bottom: 0.5rem;"><h2>🌐 Database Analysis</h2><p>Connect to a SQL database and ask business questions in plain English.</p></div>', unsafe_allow_html=True)
        if st.button("Go to Database Connector", key="select_db", use_container_width=True):
            st.session_state.app_mode = "db_analysis"
            safe_rerun()
            
    with col2:
        # Retain custom style with explicit button for reliability
        st.markdown('<div class="mode-card" style="margin-bottom: 0.5rem;"><h2>📄 Document Analysis </h2><p>Upload documents (PDFs, CSVs, etc.) and query them using AI.</p></div>', unsafe_allow_html=True)
        if st.button("Go to Document Analysis Hub", key="select_doc", use_container_width=True):
            st.session_state.app_mode = "doc_analysis"
            safe_rerun()

def display_database_login():
    # --- ADD THIS BUTTON HERE ---
    st.sidebar.button("🏠 Go to Home Page / Disconnect", key="db_disconnect", 
                   on_click=st.session_state.clear)

        
    st.title("🔗 Database Connector")
    st.markdown("Use this interface to securely connect to your database for AI-powered analysis.")

    # --- START: FIX BY USING st.form ---
    with st.form(key="db_login_form", clear_on_submit=False):
        
        # Input Fields
        db_type = st.selectbox("Database Type", ["Snowflake", "MySQL", "PostgreSQL", "SQL Server"], key="f_db_type")
        
        # Determine inputs based on type
        if db_type == "Snowflake":
            # Snowflake specific inputs
            account = st.text_input("Account (e.g., ab12345.eu-west-1)", key="f_account")
            warehouse = st.text_input("Warehouse", key="f_warehouse")
            database = st.text_input("Database Name", key="f_database")
            schema = st.text_input("Schema (Optional)", value=st.session_state.f_schema, key="f_schema")
            user = st.text_input("User", key="f_user_sf")
            password = st.text_input("Password", type="password", key="f_password_sf")
            port = 443 
            dbname = database 
            host = None
        else:
            # General inputs
            host = st.text_input("Host", key="f_host")
            # Set default port based on selection
            default_port = "3306" if db_type == "MySQL" else "5432" if db_type == "PostgreSQL" else "1433"
            port = st.text_input("Port", value=default_port, key="f_port")
            dbname = st.text_input("Database Name (or SID/Initial Catalog)", key="f_dbname")
            user = st.text_input("User", key="f_user_gen")
            password = st.text_input("Password", type="password", key="f_password_gen")
            # Set values for connection logic
            account = None 
            warehouse = None
            database = None
            schema = None

        # The new submit button that only triggers a rerun ON CLICK
        submitted = st.form_submit_button("Connect to Database")
    # --- END: FIX BY USING st.form ---

   
    
    # Logic runs only on submit, preventing intermediate reruns
    if submitted:
        # Define credentials based on the form inputs
        credentials = {
            'user': user,
            'password': password,
        }
        
        if db_type == "Snowflake":
            credentials.update({'account': account, 'warehouse': warehouse, 'database': database})
        else:
            credentials.update({'host': host, 'port': port, 'dbname': dbname})

        with st.spinner(f"Connecting to {db_type}..."):
            conn = connect_to_database(db_type, credentials)

        if conn:
            st.session_state.logged_in = True
            st.session_state.db_conn = conn
            st.session_state.db_type = db_type
            
            # Fetch and store schema immediately after successful connection
            cols, fks = get_db_schema(conn, db_type)
            st.session_state.db_schema_columns = cols
            st.session_state.db_schema_fks = fks
            
            st.success(f"Connection successful to {db_type}! Schema fetched.")
            safe_rerun()
        # else: error is handled inside connect_to_database

def display_database_analysis_page():
    st.title(f"📊  Analysis")
    
    # NEW: Direct swap button
    if st.sidebar.button("📄 Go to Document Analysis", key="swap_to_doc"):
        st.session_state.app_mode = "doc_analysis"
        safe_rerun()
    
    # EXISTING: Disconnect/Home button (modified text)
    st.sidebar.button("🏠 Go to Home Page", key="doc_disconnect", on_click=st.session_state.clear)
    
    st.markdown('***')
    
    # --- RAG System Management ---
    with st.expander("🛠️ RAG System and Schema"):
        st.subheader("Document Context for Enhanced SQL")
        
        # --- START: FIX BY USING st.form for RAG Upload ---
        with st.form(key="db_rag_form", clear_on_submit=True):
            uploaded_files = st.file_uploader(
                "Upload documents for context (PDF, DOCX, CSV, XLSX, JSON, TXT)", 
                type=['pdf', 'docx', 'csv', 'xlsx', 'json', 'txt'], 
                accept_multiple_files=True,
                key="db_rag_uploader"
            )
            upload_submitted = st.form_submit_button("Upload & Index Document(s)")
            
        if upload_submitted and uploaded_files:
            for file in uploaded_files:
                with st.spinner(f"Processing {file.name}..."):
                    result = rag_system.process_document(file)
                    if result['status'] == 'success':
                        st.session_state.uploaded_documents.append(file.name)
                        st.success(f"Indexed {result['chunks_added']} chunks from {file.name}")
                    else:
                        st.error(f"Failed to process {file.name}: {result['message']}")
            st.session_state.rag_clear_success = False
        # --- END: FIX BY USING st.form for RAG Upload ---

        # Vector DB Stats and Clear Button (use on_click to avoid form overhead)
        stats = rag_system.get_collection_stats()
        col1, col2 = st.columns([0.7, 0.3])
        with col1:
            if 'total_chunks' in stats and stats['total_chunks'] > 0:
                st.info(f"Vector Database contains **{stats['total_chunks']}** indexed chunks.")
                if st.session_state.rag_clear_success:
                    st.success("Vector database cleared successfully!")
                    st.session_state.rag_clear_success = False
            else:
                st.info("Vector Database is empty. Upload documents to provide business context.")

        with col2:
            # Clear button uses on_click to execute the action, rerunning the page
            st.button("Clear Vector Database", key="clear_db_db", on_click=clear_db_callback)
            
        st.subheader("Database Schema Summary")
        if st.session_state.get('db_schema_columns'):
            schema_summary = build_schema_card(st.session_state.db_schema_columns, st.session_state.db_schema_fks)
            st.code(schema_summary, language='text')
        else:
            st.warning("Schema not loaded. Please ensure connection is active.")


    # --- Main Query Input & Analysis Generation ---
    st.markdown("---")
    
    # --- START: FIX BY USING st.form for Main Analysis ---
    with st.form(key="db_analysis_form"):
        question = st.text_area("Ask question...", value=st.session_state.last_question or "", key="db_question_input_f")
        submitted = st.form_submit_button("Generate Analysis")
    # --- END: FIX BY USING st.form for Main Analysis ---
    
    if submitted and question:
        reset_analysis_state() # Clear previous results
        st.session_state.last_question = question
        
        conn = st.session_state.db_conn
        db_type = st.session_state.db_type
        
        # 1. Select relevant tables and build schema card
        relevant_tables = pick_relevant_tables(st.session_state.db_schema_columns, question)
        schema_card = build_schema_card(st.session_state.db_schema_columns, st.session_state.db_schema_fks, relevant_tables)

        # 2. Get RAG Context
        with st.spinner("Searching documents for business context..."):
            rag_context = rag_system.get_relevant_context(question)
        
        # 3. Generate SQL
        with st.spinner("Generating SQL query with Gemini..."):
            sql_query = generate_sql_with_gemini(question, schema_card, db_type, rag_context)
            st.session_state.sql_query = sql_query
            
        if not is_safe_select(sql_query):
            st.error("Generated query is unsafe or empty. Analysis halted.")
            st.session_state.sql_query = sql_query # Keep unsafe query for debug
            return

        # 4. Execute SQL
        with st.spinner("Executing SQL query..."):
            try:
                df = pd.read_sql(sql_query, conn)
                st.session_state.query_result = df
            except Exception as e:
                st.error(f"Query execution failed: {e}")
                st.code(sql_query, language="sql")
                return

        # 5. Generate Insights and Visualization Code
        if not df.empty:
            with st.spinner("Generating business insights and visualization code..."):
                st.session_state.insight = generate_insight_with_gemini(df, question, rag_context)
                st.session_state.viz_code = generate_visualization_code(df, question)
        
        safe_rerun() # Rerun to display results
    
    
    # --- Display Results ---
    df = st.session_state.query_result
    insight = st.session_state.insight
    viz_code = st.session_state.viz_code

    if insight:
        st.markdown("## ✨ AI Generated Insight")
        st.markdown(insight, unsafe_allow_html=True)
        
        
        # --- END ADDED ---
        st.markdown("---")
        st.markdown("## 📈 Visualizations")
        # Execute the visualization code
        if viz_code and not viz_code.startswith('#'):
            try:
                # Local scope for execution
                loc = {'df': df, 'plt': plt, 'sns': sns}
                exec(viz_code, globals(), loc)
                
                # Retrieve the figure object
                if 'fig' in loc:
                    st.pyplot(loc['fig'])
                else:
                    st.error("Visualization code executed, but no `fig` object was found.")
            except Exception as e:
                st.error(f"Error generating visualization: {e}")
                
        else:
            st.info("Visualization code not generated (e.g., due to limited data columns).")
            
        # --- ADDED: Download Buttons for Analysis ---
    if insight:
        pdf_bytes_doc = create_html_pdf(insight, st.session_state.last_question)
        txt_bytes_doc = insight.encode('utf-8')

        st.markdown("### Download Analysis")
        col_pdf_doc, col_txt_doc = st.columns(2)

        with col_pdf_doc:
            st.download_button(
                label="⬇️ Download Analysis as PDF Report",
                data=pdf_bytes_doc,
                file_name=f"doc_analysis_report_{time.strftime('%Y%m%d%H%M%S')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        with col_txt_doc:
            st.download_button(
                label="⬇️ Download Analysis as Text",
                data=txt_bytes_doc,
                file_name=f"doc_analysis_report_{time.strftime('%Y%m%d%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )

    # --- Email Results Section ---
    if insight:
        st.markdown("---")
        st.markdown("## 📧 Share Report")

        # Check if email already sent
        if st.session_state.email_sent:
            st.info("📧 Email sent successfully! You can send another email if needed.")
            # Use on_click for simple state reset
            st.button("🔄 Send Another Email", key="send_another_email", on_click=reset_email_state)
        else:
            # --- START: FIX BY USING st.form for Email ---
            with st.form(key="db_email_form"):
                col_email, col_format = st.columns([0.6, 0.4])
                with col_email:
                    recipient_email = st.text_input("Recipient Email", value=st.session_state.recipient_email or "", key="f_db_email")
                with col_format:
                    email_format = st.selectbox("Attachment Format", ["PDF Report", "CSV Data", "Excel Data"], key="f_db_email_format")
                    
                email_message = st.text_area("Email Message (Optional)", value=st.session_state.email_message or "Please find the attached data analysis report.", key="f_db_email_message")
                
                # The submit button
                email_submitted = st.form_submit_button("Send Analysis via Email")

            if email_submitted:
                # Store inputs in session state
                st.session_state.recipient_email = recipient_email
                st.session_state.email_message = email_message
                st.session_state.email_format = email_format
                
                if not recipient_email or "@" not in recipient_email:
                    st.error("Please enter a valid recipient email address.")
                    safe_rerun()
                else:
                    pdf_bytes_email, csv_bytes, excel_bytes = None, None, None
                    
                    if email_format == "PDF Report":
                        pdf_bytes_email = create_html_pdf(insight, st.session_state.last_question)
                    elif email_format == "CSV Data":
                        csv_bytes = df.to_csv(index=False).encode('utf-8')
                    elif email_format == "Excel Data":
                        excel_buffer = BytesIO()
                        df.to_excel(excel_buffer, index=False)
                        excel_bytes = excel_buffer.getvalue()
                        
                    with st.spinner(f"Sending email to {recipient_email}..."):
                        success = send_analysis_email(
                            recipient_email=recipient_email,
                            subject=f"AI Data Analysis Report: {st.session_state.last_question}",
                            message=email_message.replace(chr(10), '<br>'),
                            pdf_bytes=pdf_bytes_email,
                            csv_bytes=csv_bytes,
                            excel_bytes=excel_bytes
                        )
                        
                    if success:
                        st.session_state.email_sent = True
                        st.success("Email sent successfully!")
                    else:
                        st.error("Failed to send email. Check API credentials and logs.")
                    
                    safe_rerun()
            # --- END: FIX BY USING st.form for Email ---


    # --- Technical Details ---
    with st.expander("🔍 Technical Details"):
        st.subheader("SQL Query")
        st.code(st.session_state.sql_query or "No SQL query available.", language="sql")
        
        st.subheader("Data Preview")
        if df is not None and not df.empty:
            st.dataframe(df, use_container_width=True)
            
            # --- Download Button for CSV Data (kept from previous turn) ---
            csv_data = df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="⬇️ Download Data as CSV",
                data=csv_data,
                file_name=f"analysis_data_{time.strftime('%Y%m%d%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )
            # --- END DOWNLOAD ---
        elif df is not None:
            st.info("Query executed successfully, but returned no data to preview.")


def display_document_analysis_page():
    st.title("📄 Document Analysis")
    
    # NEW: Direct swap button
    if st.sidebar.button("🌐 Go to Database Analysis", key="swap_to_db"):
        st.session_state.app_mode = "db_analysis"
        safe_rerun()
        
    # EXISTING: Home button (modified text)
    st.sidebar.button("🏠 Go to Home Page / Disconnect", key="db_disconnect", 
                   on_click=st.session_state.clear)
    
    # --- Vector Database Management ---
    with st.expander(" Upload Documents For Analysis"):
        st.subheader("Document Indexing")
        
        # --- START: FIX BY USING st.form for RAG Upload ---
        with st.form(key="doc_rag_form", clear_on_submit=True):
            uploaded_files = st.file_uploader(
                "Upload documents for analysis (PDF, DOCX, CSV, XLSX, JSON, TXT)",
                type=['pdf', 'docx', 'csv', 'xlsx', 'json', 'txt'],
                accept_multiple_files=True,
                key="doc_rag_uploader"
            )
            upload_submitted = st.form_submit_button("Upload & Index Document(s)")

        if upload_submitted and uploaded_files:
            for file in uploaded_files:
                with st.spinner(f"Processing {file.name}..."):
                    result = rag_system.process_document(file)
                    if result['status'] == 'success':
                        st.session_state.uploaded_documents.append(file.name)
                        st.success(f"Indexed {result['chunks_added']} chunks from {file.name}")
                    else:
                        st.error(f"Failed to process {file.name}: {result['message']}")
            st.session_state.rag_clear_success = False

        # --- END: FIX BY USING st.form for RAG Upload ---

        # Vector DB Stats and Clear Button
        stats = rag_system.get_collection_stats()
        col1, col2 = st.columns([0.7, 0.3])
        with col1:
            if 'total_chunks' in stats and stats['total_chunks'] > 0:
                st.info(f"Vector Database contains **{stats['total_chunks']}** indexed chunks from {len(set(st.session_state.uploaded_documents))} files.")
                if st.session_state.rag_clear_success:
                    st.success("Vector database cleared successfully!")
                    st.session_state.rag_clear_success = False
            else:
                st.info("Vector Database is empty. Upload documents to begin analysis.")

        with col2:
            st.button("Clear Vector Database", key="clear_db_doc", on_click=clear_db_callback)

    st.markdown("---")
    
    # --- Main Analysis Input & Generation ---
    # --- START: FIX BY USING st.form for Main Analysis ---
    with st.form(key="doc_analysis_form"):
        question = st.text_area("Ask your question about the documents...", value=st.session_state.doc_analysis_results['question'] or "", key="doc_question_input_f")
        submitted = st.form_submit_button("Analyze Documents")

    if submitted and question:
        st.session_state.doc_analysis_results['question'] = question
        st.session_state.email_sent = False
        
        with st.spinner("Searching for relevant document context..."):
            rag_context = rag_system.get_relevant_context(question)
        
        with st.spinner("Generating document analysis with Gemini..."):
            analysis = generate_document_analysis_with_rag(question, rag_context)
            
            st.session_state.doc_analysis_results.update({
                'analysis': analysis,
                'rag_context': rag_context,
                'analysis_type': 'document'
            })
        
        safe_rerun() # Rerun to display results
    # --- END: FIX BY USING st.form for Main Analysis ---

    # --- Display Results ---
    analysis = st.session_state.doc_analysis_results.get('analysis')
    question = st.session_state.doc_analysis_results.get('question')
    
    
    if analysis:
        st.markdown(f"## 🔍 Analysis for: *{question}*")
        st.markdown(analysis, unsafe_allow_html=True)
        
        # --- ADDED: Download Buttons for Analysis ---
        pdf_bytes_doc = create_html_pdf(analysis, question)
        txt_bytes_doc = analysis.encode('utf-8')
        
        st.markdown("### Download Analysis")
        col_pdf_doc, col_txt_doc = st.columns(2)

        with col_pdf_doc:
            st.download_button(
                label="⬇️ Download Analysis as PDF Report",
                data=pdf_bytes_doc,
                file_name=f"doc_analysis_report_{time.strftime('%Y%m%d%H%M%S')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
        with col_txt_doc:
            st.download_button(
                label="⬇️ Download Analysis as TXT File",
                data=txt_bytes_doc,
                file_name=f"doc_analysis_report_{time.strftime('%Y%m%d%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        # --- END ADDED ---
        
        if analysis:
            st.markdown("---")
            st.markdown("## 📧 Share Report")
    
            # Check if email already sent
            if st.session_state.email_sent:
                st.info("📧 Email sent successfully! You can send another email if needed.")
                # Use on_click for simple state reset
                st.button("🔄 Send Another Email", key="send_another_email", on_click=reset_email_state)
            else:
                # --- START: FIX BY USING st.form for Email ---
                with st.form(key="db_email_form"):
                    col_email, col_format = st.columns([0.6, 0.4])
                    with col_email:
                        recipient_email = st.text_input("Recipient Email", value=st.session_state.recipient_email or "", key="f_db_email")
                    with col_format:
                        email_format = st.selectbox("Attachment Format", ["PDF Report"], key="f_db_email_format")
                        
                    email_message = st.text_area("Email Message (Optional)", value=st.session_state.email_message or "Please find the attached data analysis report.", key="f_db_email_message")
                    
                    # The submit button
                    email_submitted = st.form_submit_button("Send Analysis via Email")
    
                if email_submitted:
                    # Store inputs in session state
                    st.session_state.recipient_email = recipient_email
                    st.session_state.email_message = email_message
                    st.session_state.email_format = email_format
                    
                    if not recipient_email or "@" not in recipient_email:
                        st.error("Please enter a valid recipient email address.")
                        safe_rerun()
                    else:
                        pdf_bytes_email, csv_bytes, excel_bytes = None, None, None
                        
                        if email_format == "PDF Report":
                            pdf_bytes_email = create_html_pdf(analysis, question)
                        elif email_format == "CSV Data":
                            csv_bytes = df.to_csv(index=False).encode('utf-8')
                        elif email_format == "Excel Data":
                            excel_buffer = BytesIO()
                            df.to_excel(excel_buffer, index=False)
                            excel_bytes = excel_buffer.getvalue()
                            
                        with st.spinner(f"Sending email to {recipient_email}..."):
                            success = send_analysis_email(
                                recipient_email=recipient_email,
                                subject=f"AI Data Analysis Report: {st.session_state.last_question}",
                                message=email_message.replace(chr(10), '<br>'),
                                pdf_bytes=pdf_bytes_email,
                                csv_bytes=csv_bytes,
                                excel_bytes=excel_bytes
                            )
                            
                        if success:
                            st.session_state.email_sent = True
                            st.success("Email sent successfully!")
                        else:
                            st.error("Failed to send email. Check API credentials and logs.")
                        
                        safe_rerun()
            # --- END: FIX BY USING st.form for Email ---

    
    # --- Technical Details ---
    with st.expander("🔍 RAG Context Details"):
        st.subheader("Relevant Document Excerpts")
        st.code(st.session_state.doc_analysis_results.get('rag_context') or "No relevant document context found.", language="text")

# --------------------------------------------------------------------------------
# --- Main Application Flow --
# --------------------------------------------------------------------------------


def main():
    # Initialize session management FIRST
    manage_session_persistence()
    init_session_state()
    
    # Mode selection page
    if not st.session_state.app_mode:
        display_mode_selection()
    
    # Database analysis flow
    elif st.session_state.app_mode == "db_analysis" and not st.session_state.logged_in:
        display_database_login()
    
    elif st.session_state.app_mode == "db_analysis" and st.session_state.logged_in:
        display_database_analysis_page()
    
    # Document analysis flow
    elif st.session_state.app_mode == "doc_analysis":
        display_document_analysis_page()
        
    
if __name__ == "__main__":

    main()











